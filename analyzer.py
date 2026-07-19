"""
Core analysis engine for the resume analyzer. Runs entirely locally with
regex/scikit-learn, so the app works with zero API keys.
"""

import os
import re
import difflib
from datetime import datetime

from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

from flashtext import KeywordProcessor

import ml_models
from ml_models import (
    sbert_similarity,
    extract_entities_spacy,
    build_ats_feature_vector,
    predict_ats_score,
)

def get_st_model():
    # Backwards-compatible alias; real implementation lives in ml_models.
    return ml_models.get_sbert_model()

SKILL_DB = {
    "languages": [
        "python", "java", "javascript", "typescript", "c++", "c#", "go", "golang",
        "rust", "ruby", "php", "swift", "kotlin", "scala", "r", "matlab", "sql",
        "html", "css", "bash", "shell", "perl", "objective-c", "dart",
    ],
    "frameworks": [
        "react", "angular", "vue", "next.js", "django", "flask", "fastapi",
        "spring", "spring boot", "express", "node.js", "rails", ".net", "laravel",
        "tensorflow", "pytorch", "keras", "scikit-learn", "pandas", "numpy",
        "bootstrap", "tailwind", "jquery", "redux", "graphql",
    ],
    "tools": [
        "git", "docker", "kubernetes", "jenkins", "aws", "azure", "gcp",
        "terraform", "ansible", "linux", "jira", "confluence", "figma",
        "postman", "webpack", "ci/cd", "github actions", "nginx", "redis",
        "mongodb", "postgresql", "mysql", "elasticsearch", "kafka", "spark",
        "hadoop", "tableau", "power bi", "excel", "airflow", "grafana",
    ],
    "soft_skills": [
        "leadership", "communication", "teamwork", "problem solving",
        "project management", "time management", "collaboration",
        "critical thinking", "adaptability", "mentoring", "public speaking",
        "negotiation", "agile", "scrum", "stakeholder management",
    ],
}
ALL_SKILLS = sorted({s for group in SKILL_DB.values() for s in group}, key=len, reverse=True)

ACTION_VERBS = {
    "achieved", "accelerated", "built", "created", "delivered", "designed",
    "developed", "drove", "engineered", "established", "executed", "expanded",
    "generated", "improved", "implemented", "increased", "initiated", "launched",
    "led", "managed", "optimized", "orchestrated", "pioneered", "reduced",
    "resolved", "spearheaded", "streamlined", "transformed", "boosted",
    "automated", "architected", "negotiated", "mentored", "scaled",
    "deployed", "served", "shipped", "maintained", "integrated", "authored",
    "migrated", "configured", "monitored", "processed", "analyzed", "tested",
    "debugged", "refactored", "wrote", "trained", "containerized",
    "provisioned", "presented", "published", "collaborated", "coordinated",
    "supported", "enhanced", "upgraded", "consolidated", "standardized",
    "modernized", "documented", "administered", "diagnosed", "instrumented",
}

WEAK_PHRASES = [
    "hardworking", "team player", "results-driven", "results driven",
    "detail-oriented", "detail oriented", "go-getter", "self-starter",
    "think outside the box", "synergy", "dynamic individual", "people person",
    "hard worker", "responsible for", "duties included",
]

VERB_HINTS = [
    (r"\b(manag|oversaw|oversee|supervis)\w*\b", "Managed"),
    (r"\b(built|build|develop|creat)\w*\b", "Built"),
    (r"\b(design)\w*\b", "Designed"),
    (r"\b(improv|optimi[sz]e|enhanc)\w*\b", "Optimized"),
    (r"\b(lead|led)\w*\b", "Led"),
    (r"\b(analy[sz]e|research)\w*\b", "Analyzed"),
    (r"\b(automat)\w*\b", "Automated"),
    (r"\b(reduc|decreas|cut|lower)\w*\b", "Reduced"),
    (r"\b(increas|grow|grew|boost|scal)\w*\b", "Increased"),
    (r"\b(launch|deploy|releas|ship)\w*\b", "Launched"),
    (r"\b(coordinat|organiz)\w*\b", "Coordinated"),
    (r"\b(support|assist|help)\w*\b", "Supported"),
]
DEFAULT_VERB = "Drove"

WEAK_OPENERS_RE = re.compile(
    r"^(responsible for|duties included|worked on|helped with|helped to|"
    r"assisted with|tasked with|in charge of)\s*",
    re.I,
)

DOMAINS = {
    "Frontend Development": {"react", "vue", "angular", "html", "css", "javascript", "typescript", "tailwind css", "next.js", "bootstrap", "frontend", "ui", "ux", "figma"},
    "Backend Development": {"node.js", "python", "java", "django", "flask", "spring", "c#", "ruby", "php", "go", "golang", "postgresql", "mysql", "mongodb", "redis", "elasticsearch", "kafka", "backend", "api"},
    "Data Science & AI": {"python", "machine learning", "artificial intelligence", "deep learning", "tensorflow", "pytorch", "scikit-learn", "pandas", "numpy", "data science", "nlp", "sql", "r", "spark", "hadoop"},
    "DevOps & Cloud": {"aws", "gcp", "azure", "docker", "kubernetes", "jenkins", "terraform", "ansible", "linux", "ci/cd", "github actions", "nginx", "bash", "shell", "devops"},
    "Mobile Development": {"swift", "kotlin", "react native", "flutter", "dart", "objective-c", "mobile", "ios", "android"},
}

# Matches real-world header variants ("Key Projects", "3. Projects", etc.)
# by allowing a leading qualifier word and decorative bullets/dashes, while
# still requiring the header be essentially standalone.
_HEAD_LEAD = r"^\s*[-=~#>*•\u2022\u2500\u2501\u2504\u2508│┃]*\s*(\d+[\.\)]\s*)?"
_HEAD_TAIL = r"\s*[-=~#>*•\u2022\u2500\u2501\u2504\u2508│┃:|]*\s*$"

SECTION_HEADER_PATTERNS = [
    ("summary", _HEAD_LEAD + r"(professional\s+|career\s+)?(summary|objective|profile|about\s*me)" + _HEAD_TAIL),
    ("experience", _HEAD_LEAD + r"((professional|relevant|work)\s+)?(experience|employment(\s+history)?|work\s+history|internships?)" + _HEAD_TAIL),
    ("education", _HEAD_LEAD + r"(education(al)?(\s+background)?|academics?|academic\s+background|qualifications)" + _HEAD_TAIL),
    ("skills", _HEAD_LEAD + r"((technical|core|key|professional)\s+)?(skills|competenc(y|ies)|technologies)" + _HEAD_TAIL),
    ("projects", _HEAD_LEAD + r"((key|notable|personal|academic|major|github|side|featured|selected)\s+)?projects?" + _HEAD_TAIL),
]

SECTION_PATTERNS = {
    "contact": r"(email|phone|linkedin|@)",
    "summary": r"(summary|objective|profile|professional summary)\b",
    "experience": r"(experience|employment|work history|internships?)\b",
    "education": r"(education|academic|academics)\b",
    "skills": r"(skills|technical skills|competencies|technologies)\b",
    "projects": r"(projects?)\b",
}

STOPWORDS = set("""
a an the and or of to in on for with as at by from is are was were be been
being this that these those it its your you i we our their his her they he
she will shall can could should would may might must not no nor so than then
""".split())

# Generic job-description boilerplate that shouldn't surface as a "missing keyword"
JD_FILLER_WORDS = set("""
need needs needed looking look strong required requirement requirements
experience experienced years year ability abilities skill skills work
working team teams role roles job jobs candidate candidates plus preferred
minimum environment including include includes excellent proven demonstrated
knowledge understanding familiarity responsible responsibilities
opportunity opportunities company companies join great good etc
""".split())

def extract_text(filepath, filename):
    ext = filename.rsplit(".", 1)[-1].lower()
    if ext == "pdf":
        try:
            import pdfplumber
        except ImportError:
            raise ValueError(
                "PDF support isn't installed on this server. Run "
                "'pip install -r requirements.txt' (or just 'pip install pdfplumber') "
                "in the same Python environment used to start app.py, then restart the server."
            )
        text = []
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                text.append(page.extract_text() or "")
        return "\n".join(text)
    elif ext == "docx":
        try:
            import docx
        except ImportError:
            raise ValueError(
                "DOCX support isn't installed on this server. Run "
                "'pip install -r requirements.txt' (or just 'pip install python-docx') "
                "in the same Python environment used to start app.py, then restart the server."
            )
        doc = docx.Document(filepath)
        return "\n".join(p.text for p in doc.paragraphs)
    elif ext == "txt":
        with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    else:
        raise ValueError(f"Unsupported file type: .{ext}")

def extract_contact_info(text):
    email = re.search(r"[\w.+-]+@[\w-]+\.[\w.-]+", text)
    # Try a labeled number ("Phone: ..."), then a separator-delimited number,
    # then a bare 10-13 digit run not touching other digits.
    phone = None
    labeled = re.search(
        r"(?:phone|mobile|cell|tel(?:ephone)?|contact(?:\s*(?:no|number))?|whatsapp)\s*[:.\-]?\s*"
        r"(\+?\d{1,3}[\s.-]?(?:\(\d{2,4}\)[\s.-]?)?\d[\d\s.-]{6,13}\d)",
        text,
        re.I,
    )
    if labeled:
        phone = labeled.group(1).strip()
    else:
        separated = re.search(
            r"(?<!\d)(\+?\d{1,3}[\s.-]?)?(\(\d{3}\)|\d{3})[\s.-]\d{3}[\s.-]?\d{4}(?!\d)",
            text,
        )
        if separated:
            phone = separated.group(0)
        else:
            bare = re.search(r"(?<!\d)(\+\d{1,3}[\s-]?)?\d{10,13}(?!\d)", text)
            if bare:
                phone = bare.group(0)
    linkedin = re.search(r"linkedin\.com/in/[\w-]+", text, re.I)
    return {
        "email": email.group(0) if email else None,
        "phone": phone,
        "linkedin": linkedin.group(0) if linkedin else None,
    }
def detect_sections(text):
    lower = text.lower()
    return {name: bool(re.search(pattern, lower)) for name, pattern in SECTION_PATTERNS.items()}

# FlashText Skill Extractor Setup
SKILL_SYNONYMS = {
    "React": ["react", "react.js", "reactjs"],
    "Python": ["python", "python3"],
    "Node.js": ["node.js", "node js", "nodejs", "node"],
    "Machine Learning": ["machine learning", "ml"],
    "Artificial Intelligence": ["artificial intelligence", "ai"],
    "C++": ["c++", "cpp"],
    "C#": ["c#", "csharp"],
    "Vue.js": ["vue", "vue.js", "vuejs"],
    "Next.js": ["next.js", "nextjs"],
    "Tailwind CSS": ["tailwind", "tailwindcss", "tailwind css"],
    "PostgreSQL": ["postgresql", "postgres"],
    "MongoDB": ["mongodb", "mongo"],
    "Amazon Web Services (AWS)": ["aws", "amazon web services"],
    "Google Cloud Platform (GCP)": ["gcp", "google cloud", "google cloud platform"],
    "Natural Language Processing": ["nlp", "natural language processing"],
    "Deep Learning": ["deep learning", "dl"],
    "Data Science": ["data science", "ds"]
}

keyword_processor = KeywordProcessor(case_sensitive=False)
for standard_name, synonyms in SKILL_SYNONYMS.items():
    keyword_processor.add_keywords_from_dict({standard_name: synonyms})
for category, skills in SKILL_DB.items():
    for skill in skills:
        found = False
        for standard, syns in SKILL_SYNONYMS.items():
            if skill.lower() in syns or skill.lower() == standard.lower():
                found = True
                break
        if not found:
            keyword_processor.add_keyword(skill, skill.title())
def extract_skills(text):
    return set(keyword_processor.extract_keywords(text))
def detect_domain(skills_found):
    skills_lower = {s.lower() for s in skills_found}
    best_domain = "General Software Engineering"
    max_overlap = 0
    for domain, keywords in DOMAINS.items():
        overlap = len(skills_lower.intersection(keywords))
        if overlap > max_overlap and overlap >= 2:
            max_overlap = overlap
            best_domain = domain
    missing_core_skills = []
    if best_domain != "General Software Engineering":
        core = DOMAINS[best_domain]
        missing_core_skills = list(core - skills_lower)[:5]
    return {"domain": best_domain, "missing_core_skills": [s.title() for s in missing_core_skills]}
def parse_date(date_str):
    date_str = re.sub(r'[^a-zA-Z0-9]', ' ', date_str).strip().lower()
    if date_str in ["present", "current", "now"]:
        return datetime.now()
    year_match = re.search(r'\b(19|20)\d{2}\b', date_str)
    if not year_match:
        return None
    year = int(year_match.group(0))
    month = 1
    months = ["jan", "feb", "mar", "apr", "may", "jun", "jul", "aug", "sep", "oct", "nov", "dec"]
    for i, m in enumerate(months):
        if m in date_str:
            month = i + 1
            break
    try:
        return datetime(year, month, 1)
    except:
        return None
def calculate_experience_duration(exp_text):
    date_ranges = re.findall(r'((?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)?[a-z]*[\s\d]*\d{4})\s*(?:-|to|–)\s*((?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)?[a-z]*[\s\d]*\d{4}|present|current|now)', exp_text, re.I)
    total_months = 0
    for start_str, end_str in date_ranges:
        start_date = parse_date(start_str)
        end_date = parse_date(end_str)
        if start_date and end_date and end_date >= start_date:
            months = (end_date.year - start_date.year) * 12 + (end_date.month - start_date.month)
            if months < 0: months = 0
            if months > 120: months = 120
            total_months += months
    return {"months": total_months, "years": round(total_months / 12, 1)}
def count_action_verb_bullets(text):
    # Reuses the bullet-detection heuristic from extract_bullet_lines()
    # below (resolved at call-time, so the ordering here is fine).
    bullets = extract_bullet_lines(text)
    total_bullets = 0
    strong_start = 0
    for clean in bullets:
        words = re.findall(r"[a-zA-Z']+", clean)
        if not words:
            continue
        total_bullets += 1
        if words[0].lower() in ACTION_VERBS:
            strong_start += 1
    return strong_start, total_bullets
def count_quantified_bullets(text):
    # Only counts numbers inside genuine bullet/achievement lines, not just
    # any line in the document (avoids counting dates in job-title lines).
    bullets = extract_bullet_lines(text)
    quantified = 0
    for clean in bullets:
        if re.search(r"\d", clean):
            quantified += 1
    return quantified
def find_weak_phrases(text):
    lower = text.lower()
    return [p for p in WEAK_PHRASES if p in lower]
def parse_job_description(jd_text):
    if not jd_text or not jd_text.strip():
        return None
    segments = {"Required": [], "Preferred": [], "Responsibilities": [], "Education": [], "Experience": []}
    lines = jd_text.splitlines()
    current_segment = "Required"
    for line in lines:
        lower = line.strip().lower()
        if not lower: continue
        if re.search(r'^(requirements|qualifications|must have|what you need)', lower):
            current_segment = "Required"
            continue
        elif re.search(r'^(nice to have|preferred|bonus|plus)', lower):
            current_segment = "Preferred"
            continue
        elif re.search(r'^(responsibilities|what you will do|duties)', lower):
            current_segment = "Responsibilities"
            continue
        if re.search(r'\b(bachelor|master|phd|degree)\b', lower):
            segments["Education"].append(line)
        if re.search(r'\d+[\+-]?\s*years?', lower):
            segments["Experience"].append(line)
        segments[current_segment].append(line)
    return {k: "\n".join(v) for k, v in segments.items()}
def extract_keywords_from_jd(jd_text, top_n=25):
    """Pull salient keywords out of a job description: known skills first,
    then top TF-IDF unigrams/bigrams as a fallback for anything not in the DB."""
    jd_skills = extract_skills(jd_text)
    try:
        vectorizer = TfidfVectorizer(
            stop_words="english", ngram_range=(1, 2), max_features=60
        )
        tfidf = vectorizer.fit_transform([jd_text])
        scores = tfidf.toarray()[0]
        terms = vectorizer.get_feature_names_out()
        ranked = sorted(zip(terms, scores), key=lambda x: x[1], reverse=True)
        extra_terms = [
            t for t, s in ranked
            if s > 0
            and t not in jd_skills
            # Skip TF-IDF terms that are themselves already-known skills (just
            # in a different casing/synonym form, e.g. "aws" vs the canonical
            # "Amazon Web Services (AWS)"), otherwise the same skill shows up
            # twice — once canonical, once raw — in matched/missing lists.
            and not extract_skills(t)
            and not any(w in JD_FILLER_WORDS for w in t.split())
        ][:top_n]
    except ValueError:
        extra_terms = []
    keywords = list(jd_skills) + extra_terms
    seen, ordered = set(), []
    for k in keywords:
        if k not in seen:
            seen.add(k)
            ordered.append(k)
    return ordered[:top_n]
def _keyword_in_text(keyword, text_lower):
    """Whole-word/whole-phrase containment check. Plain `keyword in text`
    substring checks cause false positives for short keywords — e.g. the JD
    keyword 'go' would "match" inside 'google' or 'going', and 'r' would
    match inside almost anything. This anchors on non-word/non +#. boundaries
    so only real, standalone occurrences count."""
    pattern = r"(?<![\w+#.])" + re.escape(keyword.lower()) + r"(?![\w+#])"
    return bool(re.search(pattern, text_lower))
def jd_match_score(resume_text, jd_text, resume_skills=None):
    # Model #1 - Sentence-BERT semantic resume <-> job description
    # similarity, with a TF-IDF fallback if the model isn't available.
    similarity = sbert_similarity(resume_text, jd_text)
    if similarity is None:
        vectorizer = TfidfVectorizer(stop_words="english")
        try:
            tfidf = vectorizer.fit_transform([resume_text, jd_text])
            similarity = cosine_similarity(tfidf[0:1], tfidf[1:2])[0][0]
        except ValueError:
            similarity = 0.0
    jd_segments = parse_job_description(jd_text)
    if jd_segments and (jd_segments["Required"] or jd_segments["Preferred"]):
        req_skills = extract_keywords_from_jd(jd_segments["Required"])
        pref_skills = extract_keywords_from_jd(jd_segments["Preferred"])
    else:
        req_skills = extract_keywords_from_jd(jd_text)
        pref_skills = []
    resume_lower = resume_text.lower()
    if resume_skills is None:
        resume_skills = set()
    req_matched = [k for k in req_skills if k in resume_skills or _keyword_in_text(k, resume_lower)]
    req_missing = [k for k in req_skills if k not in req_matched]
    pref_matched = [k for k in pref_skills if k in resume_skills or _keyword_in_text(k, resume_lower)]
    pref_missing = [k for k in pref_skills if k not in pref_matched]
    all_matched = set(req_matched + pref_matched)
    keyword_density = {}
    for k in all_matched:
        pattern = r"(?<![\w+#.])" + re.escape(k.lower()) + r"(?![\w+#])"
        count = len(re.findall(pattern, resume_lower))
        if count == 0:
            count = resume_lower.count(k.lower())
        if count == 0 and k in resume_skills:
            count = 1
        keyword_density[k] = count
    return {
        "similarity": round(similarity * 100, 1),
        "required_matched": req_matched,
        "required_missing": req_missing,
        "preferred_matched": pref_matched,
        "preferred_missing": pref_missing,
        "keyword_density": keyword_density,
        "education_requirements": jd_segments["Education"][:300] if jd_segments and jd_segments["Education"] else "",
        "experience_requirements": jd_segments["Experience"][:300] if jd_segments and jd_segments["Experience"] else ""
    }

# Lighter-weight sibling of jd_match_score(): instead of a full job
# description, compares the resume against a curated skill list for a
# plainly-typed role title (e.g. "Data Analyst").
TARGET_ROLE_SKILLS = {
    "data analyst": ["SQL", "Excel", "Python", "R", "Tableau", "Power BI",
                      "Statistics", "Data Visualization", "A/B Testing", "Google Analytics"],
    "data scientist": ["Python", "R", "SQL", "Machine Learning", "Statistics", "Pandas",
                        "NumPy", "Scikit-learn", "TensorFlow", "PyTorch", "Data Visualization",
                        "Deep Learning"],
    "business analyst": ["SQL", "Excel", "Tableau", "Power BI", "Requirements Gathering",
                          "Stakeholder Management", "Business Process", "Agile", "JIRA"],
    "marketing analyst": ["Google Analytics", "Excel", "SQL", "A/B Testing", "SEO",
                           "Data Visualization", "Power BI"],
    "software engineer": ["Python", "Java", "C++", "JavaScript", "Git", "Data Structures",
                           "Algorithms", "SQL", "REST API", "Docker", "Agile"],
    "frontend developer": ["JavaScript", "TypeScript", "React", "HTML", "CSS", "Redux",
                            "Webpack", "Responsive Design", "Git"],
    "backend developer": ["Python", "Java", "Node.js", "SQL", "REST API", "Docker",
                           "Kubernetes", "PostgreSQL", "MongoDB", "Redis", "Microservices"],
    "full stack developer": ["JavaScript", "React", "Node.js", "HTML", "CSS", "SQL",
                              "REST API", "Git", "Docker"],
    "devops engineer": ["Docker", "Kubernetes", "Terraform", "Ansible", "Jenkins", "CI/CD",
                         "AWS", "Azure", "GCP", "Linux", "GitHub Actions", "Grafana"],
    "machine learning engineer": ["Python", "TensorFlow", "PyTorch", "Scikit-learn", "Pandas",
                                  "NumPy", "SQL", "Docker", "Kubernetes", "Machine Learning",
                                  "Deep Learning", "MLOps"],
    "product manager": ["Stakeholder Management", "Agile", "Scrum", "Roadmapping",
                         "Market Research", "Data Analysis", "JIRA", "Communication", "Leadership"],
    "project manager": ["Agile", "Scrum", "JIRA", "Stakeholder Management", "Risk Management",
                         "Budgeting", "Communication", "Leadership"],
    "ux designer": ["Figma", "Sketch", "Wireframing", "Prototyping", "User Research",
                     "Usability Testing", "Adobe XD"],
    "ui designer": ["Figma", "Sketch", "Adobe XD", "Prototyping", "Design Systems", "Typography"],
    "qa engineer": ["Selenium", "Manual Testing", "Automation Testing", "Test Cases", "JIRA",
                     "SQL", "API Testing", "Cypress"],
    "cloud engineer": ["AWS", "Azure", "GCP", "Terraform", "Kubernetes", "Docker", "Linux",
                        "Networking"],
    "cybersecurity analyst": ["Network Security", "SIEM", "Penetration Testing", "Firewalls",
                              "Vulnerability Assessment", "Python", "Linux"],
}

# Alternate phrasings that resolve to a canonical role key above.
TARGET_ROLE_ALIASES = {
    "data analytics": "data analyst",
    "data analysis": "data analyst",
    "business analytics": "business analyst",
    "ml engineer": "machine learning engineer",
    "ai engineer": "machine learning engineer",
    "swe": "software engineer",
    "sde": "software engineer",
    "front end developer": "frontend developer",
    "front-end developer": "frontend developer",
    "frontend engineer": "frontend developer",
    "back end developer": "backend developer",
    "back-end developer": "backend developer",
    "backend engineer": "backend developer",
    "fullstack developer": "full stack developer",
    "full-stack developer": "full stack developer",
    "devops": "devops engineer",
    "sre": "devops engineer",
    "site reliability engineer": "devops engineer",
    "pm": "product manager",
    "ui/ux designer": "ux designer",
    "ux/ui designer": "ux designer",
    "qa": "qa engineer",
    "qa tester": "qa engineer",
    "test engineer": "qa engineer",
    "sdet": "qa engineer",
    "cloud architect": "cloud engineer",
    "security analyst": "cybersecurity analyst",
    "infosec analyst": "cybersecurity analyst",
}

def match_target_role(role_name):
    """Resolve a free-typed role title to one of the canonical keys in
    TARGET_ROLE_SKILLS, or None if nothing close enough is found."""
    if not role_name or not role_name.strip():
        return None
    norm = re.sub(r"[^a-z0-9+#. ]", "", role_name.strip().lower())
    norm = re.sub(r"\s+", " ", norm).strip()
    norm = TARGET_ROLE_ALIASES.get(norm, norm)
    if norm in TARGET_ROLE_SKILLS:
        return norm
    keys = list(TARGET_ROLE_SKILLS.keys())
    close = difflib.get_close_matches(norm, keys, n=1, cutoff=0.8)
    if close:
        return close[0]

    # Substring match, e.g. "senior data analyst" -> "data analyst".
    for key in keys:
        if key in norm:
            return key
    return None
def target_role_skill_gap(resume_text, resume_skills, role_name):
    """Compare the resume against the expected skill set for a plainly typed
    target role (e.g. "Data Analyst") and report matched/missing skills."""
    canonical_role = match_target_role(role_name)
    if not canonical_role:
        return {
            "role_input": role_name,
            "recognized": False,
            "matched_role": None,
            "matched_skills": [],
            "missing_skills": [],
            "match_score": 0,
        }
    expected_skills = TARGET_ROLE_SKILLS[canonical_role]
    resume_lower = resume_text.lower()
    resume_skills_lower = {s.lower() for s in (resume_skills or set())}
    matched, missing = [], []
    for skill in expected_skills:
        if skill.lower() in resume_skills_lower or _keyword_in_text(skill, resume_lower):
            matched.append(skill)
        else:
            missing.append(skill)
    match_score = round(len(matched) / len(expected_skills) * 100) if expected_skills else 0
    return {
        "role_input": role_name,
        "recognized": True,
        "matched_role": canonical_role.title(),
        "matched_skills": matched,
        "missing_skills": missing,
        "match_score": match_score,
    }
def _role_fit_blurb(match_score, exp_level):
    """Short, recruiter-style read on how ready the candidate looks for a
    role at a given match score, colored slightly by experience level."""
    if match_score >= 70:
        base = "Strong fit — most of the core skills for this role are already on the resume."
    elif match_score >= 45:
        base = "Good potential — a handful of skill gaps to close before applying."
    else:
        base = "Possible stretch role — meaningful upskilling recommended first."
    if exp_level == "Fresher" and match_score >= 45:
        base += " Worth targeting for entry-level or junior openings."
    return base
def suggest_roles(text, resume_skills, exp_level=None, top_n=5, min_score=15):
    """Recruiter-style pass: instead of the user telling us a target role
    (see target_role_skill_gap above), this looks at everything already on
    the resume and works out which roles it's actually the best fit for,
    ranked by how much of each role's expected skill set the resume covers.
    """
    resume_lower = text.lower()
    resume_skills_lower = {s.lower() for s in (resume_skills or set())}
    scored = []
    for role, expected_skills in TARGET_ROLE_SKILLS.items():
        matched = [
            s for s in expected_skills
            if s.lower() in resume_skills_lower or _keyword_in_text(s, resume_lower)
        ]
        missing = [s for s in expected_skills if s not in matched]
        match_score = round(len(matched) / len(expected_skills) * 100) if expected_skills else 0
        scored.append({
            "role": role.title(),
            "match_score": match_score,
            "matched_skills": matched,
            "missing_skills": missing,
            "blurb": _role_fit_blurb(match_score, exp_level),
        })
    scored.sort(key=lambda r: r["match_score"], reverse=True)
    qualifying = [r for r in scored if r["match_score"] >= min_score]

    # If nothing clears the bar, still surface the single best-scoring role
    # instead of an empty list.
    if not qualifying and scored:
        qualifying = scored[:1]
    return qualifying[:top_n]

def determine_experience_level(chunks):
    exp_text = chunks.get("experience", "")
    edu_text = chunks.get("education", "")
    if not exp_text.strip():
        return "Fresher"
    exp_words = len(exp_text.split())
    if exp_words < 60:
        return "Fresher"

    # Recent-grad window computed off today's date so it doesn't go stale.
    current_year = datetime.now().year
    recent_years = {str(y) for y in range(current_year - 2, current_year + 2)}
    years = [y for y in re.findall(r"(20\d{2})", edu_text) if y in recent_years]
    if years and exp_words < 150:
        return "Fresher"
    return "Experienced"
def score_projects(project_text):
    if not project_text.strip():
        return {"score": 0, "metrics": {}, "suggestions": ["No projects section found."]}
    skills_in_projects = extract_skills(project_text)
    word_count = len(project_text.split())
    strong_bullets, total_bullets = count_action_verb_bullets(project_text)
    quantified = count_quantified_bullets(project_text)
    tech_score = min(len(skills_in_projects) * 15, 100)
    impact_score = min(quantified * 20, 100)
    complexity_score = min(word_count / 1.5, 100)
    action_score = (strong_bullets / total_bullets * 100) if total_bullets else 0
    overall = round(tech_score * 0.3 + impact_score * 0.3 + complexity_score * 0.2 + action_score * 0.2)
    suggestions = []
    if tech_score < 50:
        suggestions.append("Explicitly mention the technologies, frameworks, and libraries used in your projects.")
    if impact_score < 50:
        suggestions.append("Quantify the impact of your projects (e.g., 'served 500 users', 'improved speed by 20%').")
    if action_score < 50:
        suggestions.append("Start project descriptions with strong action verbs (e.g., 'Architected', 'Developed').")
    return {
        "score": overall,
        "metrics": {
            "tech_stack_score": round(tech_score),
            "impact_score": round(impact_score),
            "complexity_score": round(complexity_score),
            "action_verbs_score": round(action_score)
        },
        "suggestions": suggestions if suggestions else ["Great project descriptions with solid technical depth and measurable impact."]
    }
def score_resume(text, jd_text=None, ats_risk=None, target_role=None):
    word_count = len(re.findall(r"\w+", text))
    sections = detect_sections(text)
    contact = extract_contact_info(text)
    skills_found = extract_skills(text)
    strong_bullets, total_bullets = count_action_verb_bullets(text)
    quantified = count_quantified_bullets(text)
    weak_phrases = find_weak_phrases(text)

    # Structure score
    structure_score = 0
    structure_checks = []
    if contact["email"]:
        structure_score += 20
        structure_checks.append(("Email found", True))
    else:
        structure_checks.append(("Email found", False))
    if contact["phone"]:
        structure_score += 10
        structure_checks.append(("Phone number found", True))
    else:
        structure_checks.append(("Phone number found", False))
    for key in ["experience", "education", "skills"]:
        if sections[key]:
            structure_score += 15
        structure_checks.append((f"'{key.capitalize()}' section present", sections[key]))
    if sections["summary"]:
        structure_score += 10
    structure_checks.append(("Summary/objective present", sections["summary"]))
    if 350 <= word_count <= 900:
        structure_score += 15
        structure_checks.append(("Resume length appropriate (350-900 words)", True))
    else:
        structure_checks.append(("Resume length appropriate (350-900 words)", False))
    ats_penalty = 0
    if ats_risk:
        if ats_risk["risk_level"] == "High":
            ats_penalty = 20
        elif ats_risk["risk_level"] == "Medium":
            ats_penalty = 10
    structure_score = min(structure_score, 100)
    structure_score = max(0, structure_score - ats_penalty)
    if ats_penalty > 0:
        structure_checks.append((f"ATS formatting penalty (-{ats_penalty} pts)", False))
    else:
        structure_checks.append(("ATS formatting clean", True))

    # Content quality score
    content_score = 0
    bullet_ratio = (strong_bullets / total_bullets) if total_bullets else 0
    content_score += round(bullet_ratio * 40)
    quant_ratio = min(quantified / max(total_bullets, 1), 1)
    content_score += round(quant_ratio * 35)
    penalty = min(len(weak_phrases) * 5, 25)
    content_score += (25 - penalty)
    content_score = max(0, min(content_score, 100))

    # Skills coverage score
    skills_score = min(len(skills_found) * 6, 100)

    # Weights differ by experience level
    chunks = split_sections(text)
    exp_level = determine_experience_level(chunks)
    exp_duration = calculate_experience_duration(chunks.get("experience", ""))
    domain_info = detect_domain(skills_found)
    project_quality = score_projects(chunks.get("projects", ""))
    if exp_level == "Fresher":
        w_struct, w_content, w_skills, w_jd = 0.25, 0.30, 0.25, 0.20
        w_struct_nj, w_content_nj, w_skills_nj = 0.30, 0.40, 0.30
    else:
        w_struct, w_content, w_skills, w_jd = 0.20, 0.40, 0.25, 0.15
        w_struct_nj, w_content_nj, w_skills_nj = 0.20, 0.50, 0.30
    result = {
        "candidate_level": exp_level,
        "experience_duration": exp_duration,
        "domain": domain_info,
        "project_quality": project_quality,
        "word_count": word_count,
        "contact": contact,
        "sections": sections,
        "structure_checks": structure_checks,
        "structure_score": structure_score,
        "content_score": content_score,
        "skills_score": skills_score,
        "skills_found": sorted(skills_found),
        "strong_action_bullets": strong_bullets,
        "total_bullets_detected": total_bullets,
        "quantified_bullets": quantified,
        "weak_phrases": weak_phrases,
    }

    # Optional: JD match
    if jd_text and jd_text.strip():
        match = jd_match_score(text, jd_text, resume_skills=skills_found)
        result["jd_match"] = match
        overall = round(
            structure_score * w_struct
            + content_score * w_content
            + skills_score * w_skills
            + match["similarity"] * w_jd
        )
    else:
        result["jd_match"] = None
        overall = round(structure_score * w_struct_nj + content_score * w_content_nj + skills_score * w_skills_nj)
    result["overall_score"] = max(0, min(overall, 100))

    # Optional: target-role skill gap
    if target_role and target_role.strip():
        result["target_role_match"] = target_role_skill_gap(text, skills_found, target_role)
    else:
        result["target_role_match"] = None

    result["suggested_roles"] = suggest_roles(text, skills_found, exp_level=exp_level)
    result["suggestions"] = build_suggestions(result)

    result["section_scores"] = section_wise_scores(text)
    result["completeness"] = completeness_score(
        text, sections, contact, skills_found, quantified, total_bullets
    )
    result["ats_contribution"] = ats_contribution_analysis(
        text, chunks=chunks, section_scores=result["section_scores"]
    )

    # spaCy: structured extraction of education/experience/certifications,
    # plus an NER-based skills cross-check. None if spaCy isn't installed.
    entities = extract_entities_spacy(
        text,
        skill_vocab=ALL_SKILLS,
        experience_text=chunks.get("experience"),
        education_text=chunks.get("education"),
    )
    result["entities"] = entities
    if entities and entities.get("ner_skills"):
        merged = sorted(set(result["skills_found"]) | set(entities["ner_skills"]))
        result["skills_found"] = merged
    return result
def build_suggestions(r):
    tips = []
    if not r["contact"]["email"]:
        tips.append("Add a professional email address near the top of your resume.")
    if not r["contact"]["phone"]:
        tips.append("Include a phone number so recruiters can reach you quickly.")
    if not r["sections"]["summary"]:
        tips.append("Add a 2-3 line summary at the top tailored to the role you want.")
    if not r["sections"]["skills"]:
        tips.append("Add a dedicated 'Skills' section so ATS systems can parse your keywords.")
    if r["word_count"] < 350:
        tips.append("Your resume looks short — add more detail on impact and responsibilities.")
    elif r["word_count"] > 900:
        tips.append("Your resume is quite long — trim it down to 1-2 pages of the most relevant content.")
    if r["total_bullets_detected"] > 0:
        ratio = r["strong_action_bullets"] / r["total_bullets_detected"]
        if ratio < 0.5:
            tips.append(
                "Start more bullet points with strong action verbs "
                "(e.g. 'Led', 'Built', 'Reduced') instead of passive phrasing."
            )
    if r["quantified_bullets"] < max(r["total_bullets_detected"] // 2, 1):
        tips.append(
            "Quantify your achievements with numbers, percentages, or dollar amounts "
            "(e.g. 'Reduced load time by 40%')."
        )
    if r["weak_phrases"]:
        tips.append(
            "Replace overused phrases like "
            + ", ".join(f"'{p}'" for p in r["weak_phrases"][:3])
            + " with specific, evidence-backed statements."
        )
    if len(r["skills_found"]) < 6:
        tips.append("List more relevant technical and soft skills explicitly — ATS software scans for exact keyword matches.")
    if r["jd_match"]:
        missing = (r["jd_match"].get("required_missing", []) + r["jd_match"].get("preferred_missing", []))[:8]
        if missing:
            tips.append(
                "The job description mentions these keywords that don't appear in your resume: "
                + ", ".join(missing)
                + ". Add the ones that genuinely apply to your experience."
            )
        if r["jd_match"]["similarity"] < 50:
            tips.append("Your resume's overall similarity to this job description is low — mirror its terminology where truthful.")
    if r.get("target_role_match"):
        trm = r["target_role_match"]
        if trm["recognized"] and trm["missing_skills"]:
            tips.append(
                f"For a {trm['matched_role']} role, your resume is missing: "
                + ", ".join(trm["missing_skills"][:8])
                + ". Add the ones that genuinely apply to your experience."
            )
        elif not trm["recognized"] and trm["role_input"]:
            tips.append(
                f"'{trm['role_input']}' wasn't recognized as a target role — try a more common title "
                "(e.g. 'Data Analyst', 'Software Engineer', 'Product Manager')."
            )
    if r.get("suggested_roles"):
        top = r["suggested_roles"][0]
        if top["match_score"] >= 45:
            tips.append(
                f"Based on your current skills, you look like a strong candidate for "
                f"{top['role']} roles ({top['match_score']}% skill match) — worth applying."
            )
    if not tips:
        tips.append("Great work — your resume covers the fundamentals well. Consider a final proofread and a tailored summary per application.")
    return tips

# Section-wise scoring

def split_sections(text):
    """Best-effort split of the resume into named chunks, based on section
    headers that appear alone on their own line. Anything before the first
    recognised header is treated as the 'header' zone (name/contact/summary
    area). This is heuristic — resumes with unusual formatting may not split
    perfectly, but it's good enough to give section-level signal."""
    lines = text.splitlines()
    sections = {}
    current = "header"
    buffer = []
    def flush():
        sections.setdefault(current, [])
        sections[current].extend(buffer)
    for line in lines:
        stripped = line.strip()
        matched = None
        for name, pattern in SECTION_HEADER_PATTERNS:
            if re.match(pattern, stripped, re.I):
                matched = name
                break
        if matched:
            flush()
            buffer = []
            current = matched
        else:
            buffer.append(line)
    flush()
    return {k: "\n".join(v) for k, v in sections.items()}
def section_wise_scores(text):
    """Score each major resume section independently (0-100), so a person
    can see exactly which part of the resume is weakest rather than just an
    overall number."""
    chunks = split_sections(text)
    scores = {}

    # Contact info can sit anywhere near the top, so pull from the whole doc.
    contact = extract_contact_info(text)
    c_score = 0
    if contact["email"]:
        c_score += 50
    if contact["phone"]:
        c_score += 30
    if contact["linkedin"]:
        c_score += 20
    scores["contact"] = {"label": "Contact info", "score": min(c_score, 100), "present": c_score > 0}

    # Summary
    summary_text = chunks.get("summary", "").strip()
    if summary_text:
        wc = len(summary_text.split())
        s_score = 100 if 20 <= wc <= 80 else (65 if wc > 0 else 0)
        scores["summary"] = {"label": "Summary / objective", "score": s_score, "present": True}
    else:
        scores["summary"] = {"label": "Summary / objective", "score": 0, "present": False}

    # Experience
    exp_text = chunks.get("experience", "")
    if exp_text.strip():
        strong, total = count_action_verb_bullets(exp_text)
        quant = count_quantified_bullets(exp_text)
        ratio_action = (strong / total) if total else 0
        ratio_quant = min(quant / max(total, 1), 1)
        e_score = round(ratio_action * 60 + ratio_quant * 40)
        scores["experience"] = {
            "label": "Experience", "score": e_score, "present": True, "bullets_detected": total,
        }
    else:
        scores["experience"] = {"label": "Experience", "score": 0, "present": False, "bullets_detected": 0}

    # Education
    edu_text = chunks.get("education", "")
    if edu_text.strip():
        has_year = bool(re.search(r"(19|20)\d{2}", edu_text))
        has_degree = bool(re.search(r"(bachelor|master|b\.?s\.?|m\.?s\.?|ph\.?d|associate|diploma|degree)", edu_text, re.I))
        ed_score = 40 + (30 if has_year else 0) + (30 if has_degree else 0)
        scores["education"] = {"label": "Education", "score": ed_score, "present": True}
    else:
        scores["education"] = {"label": "Education", "score": 0, "present": False}

    # Skills
    skills_text = chunks.get("skills", "")
    if skills_text.strip():
        found = extract_skills(skills_text)
        sk_score = min(len(found) * 10, 100)
        scores["skills"] = {"label": "Skills", "score": sk_score, "present": True, "count": len(found)}
    else:
        found = extract_skills(text)
        sk_score = min(len(found) * 5, 60) if found else 0
        scores["skills"] = {"label": "Skills", "score": sk_score, "present": bool(found), "count": len(found)}

    # Projects (optional section)
    proj_text = chunks.get("projects", "")
    if proj_text.strip():
        strong, total = count_action_verb_bullets(proj_text)
        p_score = round((strong / total) * 100) if total else 55
        scores["projects"] = {"label": "Projects", "score": p_score, "present": True}
    else:
        scores["projects"] = {"label": "Projects", "score": None, "present": False}
    return scores
def completeness_score(text, sections, contact, skills_found, quantified, total_bullets):
    has_github = bool(re.search(r"github\.com/[\w-]+", text, re.I))
    has_certifications = bool(re.search(r"^\s*(certifications?|licenses?)\s*:?\s*$", text, re.I | re.M))
    has_achievements = bool(re.search(r"^\s*(achievements?|awards?|honors?|recognitions?)\s*:?\s*$", text, re.I | re.M))
    checklist = [
        ("Summary", sections["summary"]),
        ("Skills", sections["skills"]),
        ("Experience", sections["experience"]),
        ("Projects", sections["projects"]),
        ("Education", sections["education"]),
        ("Certifications", has_certifications),
        ("Achievements", has_achievements),
        ("Contact Information", bool(contact["email"] or contact["phone"])),
        ("LinkedIn", bool(contact["linkedin"])),
        ("GitHub", has_github),
    ]
    passed = sum(1 for _, ok in checklist if ok)
    score = round(passed / len(checklist) * 100)
    missing = [label for label, ok in checklist if not ok]
    return {"score": score, "checklist": checklist, "missing": missing}

# ATS Score Contribution Analysis
#
# Breaks the resume down into the zones a real ATS parser would actually
# read (Projects, Experience, Skills, Education, Certifications, Others) and
# reports what percentage of the final ATS score each one is responsible
# for. This is purely a *reporting* layer on top of the existing scoring
# logic above (section_wise_scores / completeness_score / analyze_ats_risk)
# — it does not change how structure_score, content_score, skills_score, or
# ml_ats_score are computed anywhere else.
#
# Each section's contribution is driven by two things pulled straight from
# the resume, never hardcoded:
#   1. "weight" — how much of the resume's parseable text actually lives in
#      that section (word count share). A section an ATS never sees can't
#      meaningfully move the score.
#   2. "quality" — that section's own 0-100 score from the existing
#      section_wise_scores()/certification heuristics above.
# contribution_i = weight_i * quality_i, renormalized so all contributions
# always sum to exactly 100%.

CERT_HEADER_PATTERN = _HEAD_LEAD + r"(certifications?|licenses?|credentials?)" + _HEAD_TAIL

def extract_certifications_chunk(text):
    """Best-effort extraction of a standalone Certifications/Licenses block.
    Kept intentionally independent of split_sections() (which feeds the main
    scoring pipeline) so this reporting-only feature can never alter the
    existing section-splitting behavior used elsewhere."""
    lines = text.splitlines()
    header_patterns = SECTION_HEADER_PATTERNS + [("certifications", CERT_HEADER_PATTERN)]
    in_cert = False
    buffer = []
    for line in lines:
        stripped = line.strip()
        matched = None
        for name, pattern in header_patterns:
            if re.match(pattern, stripped, re.I):
                matched = name
                break
        if matched == "certifications":
            in_cert = True
            continue
        elif matched is not None:
            in_cert = False
            continue
        if in_cert:
            buffer.append(line)
    return "\n".join(buffer)

def score_certifications_chunk(cert_text):
    """0-100 quality score for the certifications block, in the same style
    as the other per-section scores in section_wise_scores()."""
    cert_text = (cert_text or "").strip()
    if not cert_text:
        return {"label": "Certifications", "score": 0, "present": False, "count": 0}
    lines = [l.strip() for l in cert_text.splitlines() if l.strip()]
    count = len(lines)
    score = min(40 + count * 20, 100)
    return {"label": "Certifications", "score": score, "present": True, "count": count}

def ats_contribution_analysis(text, chunks=None, section_scores=None):
    """Dynamically computes what share of the final ATS score each resume
    section is responsible for. Returns percentages that always total 100."""
    if chunks is None:
        chunks = split_sections(text)
    if section_scores is None:
        section_scores = section_wise_scores(text)

    cert_chunk = extract_certifications_chunk(text)
    cert_info = score_certifications_chunk(cert_chunk)

    total_words = max(len(re.findall(r"\w+", text)), 1)

    def wc(chunk_text):
        return len(re.findall(r"\w+", chunk_text or ""))

    experience_wc = wc(chunks.get("experience", ""))
    education_wc = wc(chunks.get("education", ""))
    skills_wc = wc(chunks.get("skills", ""))
    projects_wc = wc(chunks.get("projects", ""))
    cert_wc = wc(cert_chunk)
    header_wc = wc(chunks.get("header", ""))
    summary_wc = wc(chunks.get("summary", ""))

    # Anything not attributed to a named section (header/contact zone,
    # summary, plus any leftover text split_sections couldn't classify)
    # rolls into "Others" so all word counts always account for 100% of
    # the document, and thus the final percentages always sum to 100.
    classified_wc = experience_wc + education_wc + skills_wc + projects_wc + cert_wc
    others_wc = max(total_words - classified_wc, 0)

    contact_score = section_scores.get("contact", {}).get("score", 0) or 0
    summary_score = section_scores.get("summary", {}).get("score", 0) or 0
    # "Others" quality = blend of contact-info completeness and summary
    # quality (weighted by how much text each actually has), since those are
    # the two zones folded into it.
    if header_wc + summary_wc > 0:
        others_quality = round(
            (contact_score * header_wc + summary_score * summary_wc) / max(header_wc + summary_wc, 1)
        )
    else:
        others_quality = contact_score

    projects_score = section_scores.get("projects", {}).get("score") or 0
    experience_score = section_scores.get("experience", {}).get("score") or 0
    education_score = section_scores.get("education", {}).get("score") or 0
    skills_score = section_scores.get("skills", {}).get("score") or 0

    raw = {
        "Projects": projects_wc * projects_score,
        "Experience": experience_wc * experience_score,
        "Skills": skills_wc * skills_score,
        "Education": education_wc * education_score,
        "Certifications": cert_wc * cert_info["score"],
        "Others": others_wc * others_quality,
    }

    total_raw = sum(raw.values())
    if total_raw <= 0:
        # No section carries any weighted quality (e.g. a near-empty resume)
        # — fall back to a purely presence-based split so the chart still
        # shows something meaningful instead of all zeros.
        presence_weight = {
            "Projects": projects_wc,
            "Experience": experience_wc,
            "Skills": skills_wc,
            "Education": education_wc,
            "Certifications": cert_wc,
            "Others": max(others_wc, 1 if total_words else 0),
        }
        total_presence = sum(presence_weight.values()) or 1
        raw = presence_weight
        total_raw = total_presence

    percentages = {k: (v / total_raw * 100 if total_raw else 0) for k, v in raw.items()}

    # Round to 1 decimal place, then fix up rounding drift on the largest
    # bucket so the displayed percentages always total exactly 100.0.
    rounded = {k: round(v, 1) for k, v in percentages.items()}
    drift = round(100.0 - sum(rounded.values()), 1)
    if abs(drift) >= 0.1:
        top_key = max(rounded, key=rounded.get)
        rounded[top_key] = round(rounded[top_key] + drift, 1)

    order = ["Experience", "Skills", "Projects", "Education", "Certifications", "Others"]
    breakdown = [
        {
            "section": name,
            "percentage": max(rounded[name], 0),
            "section_score": {
                "Projects": projects_score,
                "Experience": experience_score,
                "Skills": skills_score,
                "Education": education_score,
                "Certifications": cert_info["score"],
                "Others": others_quality,
            }[name],
            "present": {
                "Projects": projects_wc > 0,
                "Experience": experience_wc > 0,
                "Skills": skills_wc > 0,
                "Education": education_wc > 0,
                "Certifications": cert_info["present"],
                "Others": (header_wc + summary_wc) > 0,
            }[name],
        }
        for name in order
    ]
    return breakdown

# ATS risk analysis

UNUSUAL_BULLET_CHARS = re.compile(r"[➤◆■●▪✦❖✔➔]")

def analyze_ats_risk(filepath, filename, text):
    """Heuristic scan for formatting choices that commonly trip up ATS
    (Applicant Tracking System) parsers: tables, embedded images, multi-column
    layouts, exotic bullet glyphs, and missing standard section headers.
    This can't see the actual visual layout/fonts, only what the parser
    itself could extract — which is exactly the same limitation a real ATS
    has, so it's a reasonable proxy."""
    ext = filename.rsplit(".", 1)[-1].lower()
    issues = []
    has_tables = False
    has_images = False
    try:
        if ext == "pdf":
            import pdfplumber
            with pdfplumber.open(filepath) as pdf:
                for page in pdf.pages:
                    try:
                        if page.find_tables():
                            has_tables = True
                    except Exception:
                        pass
                    if page.images:
                        has_images = True
        elif ext == "docx":
            import docx
            doc = docx.Document(filepath)
            if doc.tables:
                has_tables = True
            if doc.inline_shapes:
                has_images = True
    except Exception:
        pass
    if has_tables:
        issues.append({
            "issue": "Tables detected",
            "severity": "high",
            "tip": "Avoid tables for layout — many ATS parsers read table cells out of "
                   "order or skip them entirely. Use plain single-column text with clear headers instead.",
        })
    if has_images:
        issues.append({
            "issue": "Images or graphics detected",
            "severity": "medium",
            "tip": "Avoid logos, icons, headshots, or charts — ATS software cannot read "
                   "image content, and it can break parsing of nearby text.",
        })

    # Multi-column heuristic: lines with large internal gaps often indicate
    # side-by-side columns that a linear text extractor reads out of order.
    gap_lines = sum(1 for l in text.splitlines() if re.search(r"\S {4,}\S", l))
    if gap_lines > 5:
        issues.append({
            "issue": "Possible multi-column layout",
            "severity": "medium",
            "tip": "Several lines show large internal gaps, which often means a "
                   "multi-column layout. Many ATS systems read columns left-to-right "
                   "across the whole page and scramble the content. A single column is safest.",
        })
    if UNUSUAL_BULLET_CHARS.search(text):
        issues.append({
            "issue": "Non-standard bullet or icon symbols",
            "severity": "low",
            "tip": "Stick to simple bullets (- or •). Exotic symbols and icon fonts can "
                   "render as garbled characters or boxes when an ATS extracts the text.",
        })
    matched_sections = sum(
        1 for name in ["experience", "education", "skills"]
        if re.search(SECTION_PATTERNS[name], text.lower())
    )
    if matched_sections < 3:
        issues.append({
            "issue": "Non-standard or missing section headers",
            "severity": "high",
            "tip": "Use conventional headers such as 'Experience', 'Education', and "
                   "'Skills' on their own line. Creative header names (e.g. 'My Journey') "
                   "often aren't recognised by ATS section parsers.",
        })
    if not text.strip():
        issues.append({
            "issue": "No extractable text",
            "severity": "high",
            "tip": "This file appears to contain no machine-readable text (it may be a "
                   "scanned image). ATS systems cannot read scanned resumes at all — export as a text-based PDF or DOCX.",
        })
    high = sum(1 for i in issues if i["severity"] == "high")
    medium = sum(1 for i in issues if i["severity"] == "medium")
    if high:
        level = "High"
    elif medium:
        level = "Medium"
    else:
        level = "Low"

    # XGBoost: a learned 0-100 ATS-parseability score alongside the
    # rule-based risk_level/issues above.
    contact = extract_contact_info(text)
    sections = detect_sections(text)
    skills_found = extract_skills(text)
    strong_bullets, total_bullets = count_action_verb_bullets(text)
    quantified = count_quantified_bullets(text)
    num_core_sections = sum(1 for k in ["experience", "education", "skills", "summary"] if sections.get(k))
    features = build_ats_feature_vector(
        word_count=len(re.findall(r"\w+", text)),
        has_email=bool(contact["email"]),
        has_phone=bool(contact["phone"]),
        has_linkedin=bool(contact["linkedin"]),
        num_core_sections=num_core_sections,
        has_tables=has_tables,
        has_images=has_images,
        gap_line_ratio=(gap_lines / max(len(text.splitlines()), 1)),
        has_unusual_bullets=bool(UNUSUAL_BULLET_CHARS.search(text)),
        matched_standard_headers=matched_sections,
        skills_count=len(skills_found),
        bullet_count=total_bullets,
        quantified_bullet_ratio=(quantified / max(total_bullets, 1)),
    )
    ml_score = predict_ats_score(features)
    return {
        "risk_level": level,
        "issues": issues,
        "ml_ats_score": ml_score,
    }

# Bullet point rewriting

_BULLET_MARKER_RE = re.compile(r"^[•\-\*\u2022➤◆■●▪]")

def _merge_wrapped_lines(raw_lines):
    """pdfplumber/docx extraction returns one physical line per visually
    wrapped line, so a single bullet that wraps onto a second line in the
    original resume (very common for anything ending in a quantified
    result, e.g. "...delivered predictions with" / "95% accuracy.") comes
    back here as two separate lines with no bullet marker on the second
    one. Left unmerged, the first (incomplete) line gets treated as its
    own weak bullet — truncated mid-sentence — and the orphaned
    continuation line is silently dropped, which also makes a genuinely
    quantified bullet get flagged as having no measurable outcome.

    This reassembles a continuation line onto the previous line only when
    both signals agree: the previous line doesn't already end in
    sentence-ending punctuation, AND the current line starts lowercase or
    with a digit (e.g. "95%", "$10K") — the two hallmarks of a wrapped
    trailing clause rather than a genuinely new bullet/header."""
    merged = []
    for line in raw_lines:
        is_marked = bool(_BULLET_MARKER_RE.match(line))
        is_shouty_header = line.isupper() and len(line.split()) <= 6
        prev_incomplete = merged and not merged[-1].rstrip().endswith((".", ":", "!", "?"))
        looks_like_continuation = bool(re.match(r"^[a-z]", line)) or bool(re.match(r"^\d", line))
        if (
            merged
            and not is_marked
            and not is_shouty_header
            and "@" not in line and "|" not in line
            and prev_incomplete
            and looks_like_continuation
        ):
            merged[-1] = merged[-1].rstrip() + " " + line.strip()
        else:
            merged.append(line)
    return merged

def extract_bullet_lines(text):
    """Pulls out lines that actually look like resume bullet points — either
    explicitly marked (-, *, •, ...) or sentence-like lines with enough lowercase
    connective words to distinguish them from short Title Case headers like a
    job title or company name line."""
    raw_lines = [l.strip() for l in text.splitlines() if l.strip()]
    lines = _merge_wrapped_lines(raw_lines)
    bullets = []
    for l in lines:
        is_marked = bool(_BULLET_MARKER_RE.match(l))
        clean = re.sub(r"^[•\-\*\u2022➤◆■●▪]\s*", "", l).strip()
        if "@" in clean or "|" in clean:
            continue
        words = clean.split()
        if is_marked:
            if clean:
                bullets.append(clean)
            continue
        if 5 <= len(words) <= 30:
            lowercase_after_first = sum(1 for w in words[1:] if w[:1].islower())
            if lowercase_after_first >= 2:
                bullets.append(clean)
    return bullets
def is_weak_bullet(bullet):
    words = re.findall(r"[a-zA-Z']+", bullet)
    if not words:
        return True
    starts_weak = words[0].lower() not in ACTION_VERBS
    has_number = bool(re.search(r"\d", bullet))
    return starts_weak or not has_number
def suggest_verb(bullet):
    lower = bullet.lower()
    for pattern, verb in VERB_HINTS:
        if re.search(pattern, lower):
            return verb
    return DEFAULT_VERB
def rule_based_bullet_rewrite(bullet):
    words = re.findall(r"[a-zA-Z']+", bullet)
    starts_weak = not words or words[0].lower() not in ACTION_VERBS
    has_number = bool(re.search(r"\d", bullet))
    core = WEAK_OPENERS_RE.sub("", bullet).strip()
    notes = []
    if starts_weak:
        verb = suggest_verb(bullet)
        # Only lowercase when we're prepending a new leading verb.
        if core:
            core = core[0].lower() + core[1:]
        core = f"{verb} {core}".strip()
        notes.append("Now leads with a strong action verb instead of a passive/weak phrase.")
    if not has_number:
        core = core.rstrip(". ") + " — add a measurable result (e.g. 'by 25%', 'saving $10K/year', 'for 200+ users')."
        notes.append("No quantified outcome detected — add a specific number, percentage, or dollar amount.")
    if not notes:
        notes.append("Already reasonably strong — consider tightening the wording further.")
    return {"original": bullet, "suggested": core, "notes": notes}
def get_bullet_rewrites(text, max_bullets=5):
    """Returns rule-based rewrite suggestions for the weakest bullet points
    found in the resume (weak opener and/or no quantified result)."""
    chunks = split_sections(text)
    bullet_source = "\n".join([chunks.get("experience", ""), chunks.get("projects", "")]).strip()
    if not bullet_source:
        bullet_source = text  # fallback if section splitting didn't find headers
    bullets = extract_bullet_lines(bullet_source)
    weak = [b for b in bullets if is_weak_bullet(b)][:max_bullets]
    if not weak:
        return []
    return [rule_based_bullet_rewrite(b) for b in weak]

# Project Description Enhancer

def rule_based_project_enhancement(bullet):
    words = re.findall(r"[a-zA-Z']+", bullet)
    starts_weak = not words or words[0].lower() not in ACTION_VERBS
    has_number = bool(re.search(r"\d", bullet))
    # Reuses the shared skill vocabulary rather than a separate keyword list.
    mentions_tech = bool(extract_skills(bullet)) or any(
        kw in bullet.lower() for kw in ["using", "developed with", "built with", "technologies", "stack"]
    )
    core = WEAK_OPENERS_RE.sub("", bullet).strip()
    notes = []
    if starts_weak:
        verb = suggest_verb(bullet)
        if core:
            core = core[0].lower() + core[1:]
        core = f"{verb} {core}".strip()
        notes.append("Now leads with a strong action verb instead of a passive/weak phrase.")
    if not mentions_tech:
        core = core.rstrip(". ") + ", built with [add your tech stack/tools here]"
        notes.append("Explicitly name the technologies, frameworks, or tools used in this project.")
    if not has_number:
        core = core.rstrip(". ") + " — add a measurable outcome (e.g. 'used by 500+ people', 'cut load time by 40%')."
        notes.append("Quantify the project's impact or scale with a specific number, percentage, or dollar amount.")
    if not notes:
        notes.append("Good project description. Ensure it clearly states your specific contribution if it was a team project.")
    if not core.endswith("."):
        core = core.rstrip(", ") + "."
    return {"original": bullet, "suggested": core, "notes": notes}
def get_project_enhancements(text, max_bullets=3):
    chunks = split_sections(text)
    proj_text = chunks.get("projects", "").strip()
    if not proj_text:
        return []
    bullets = extract_bullet_lines(proj_text)
    if not bullets:
        lines = [l.strip() for l in proj_text.splitlines() if l.strip() and len(l.split()) > 5]
        bullets = lines[:max_bullets]
    else:
        bullets = bullets[:max_bullets]
    if not bullets:
        return []
    return [rule_based_project_enhancement(b) for b in bullets]